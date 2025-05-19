import math
import datetime
import numpy as np
from pathlib import Path
from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from task_manager import DataTaskManager
import pandas as pd
import subprocess
from Helpers.config import elpa_testing_card_col_middle,elpa_testing_card_col_elementary,elpa_testing_card_col_high
from Helpers.columns import adjust_column_widths

def create_mail_merge(juaquina_path, els_path, ela_class_path, elpa_cards_folder):
    manager = DataTaskManager()
    folders = manager.return_folders()
    
    # Read the Excel files
    juaquina_df = pd.read_excel(juaquina_path)
    # print(juaquina_df.columns)
    els_df = pd.read_excel(els_path)
    # print(els_df.columns)
    ela_class_df = pd.read_excel(ela_class_path)
    # print(ela_class_df.columns)

    # Rename primary key
    juaquina_df = juaquina_df.rename(columns={"District Student ID": "student_number"})
    juaquina_df = juaquina_df.rename(columns={"ELPA Accessibility Supports": "ELPA_Accessibility_Supports"})
    juaquina_df.drop(
        [col for col in ['Attending School', 'LEP', 'Case Manager', 'SSID', 'Last Name', 'First Name', 'Grade']
         if col in juaquina_df.columns],
        axis=1,
        inplace=True
    )

    ela_class_df = ela_class_df.rename(columns={"Student Number": "student_number"})
    ela_class_df = ela_class_df[ela_class_df['Course Name'].str.contains('ELA|Language Arts')]
    ela_class_df = ela_class_df[["student_number", "Teacher Last, First", "Expression"]]
    ela_class_df = ela_class_df.rename(columns={"Teacher Last, First": "Teacher"})

    temp_df = pd.merge(els_df, juaquina_df, on="student_number", how="outer")
    temp_df["IEP Date"] = pd.to_datetime(temp_df["IEP Date"], errors='coerce').dt.strftime('%m/%d/%Y')
    temp_df = temp_df.rename(columns={"State_StudentNumber":"SSID"})
    temp_df = temp_df.rename(columns={"IEP Date":"IEP_Date"})
    temp_df = temp_df.rename(columns={"[schools]name":"School"})
    temp_df = temp_df.rename(columns={"S_OR_STU_X.ELFg":"EL_Student"})
    temp_df["IEP_or_504"] = temp_df.apply(lambda x: "" if pd.isna(x["IEP_Date"]) else "1", axis=1)


    #Exemption handling
    cols_to_melt = ['Exempt Reading', 'Exempt Speak',
       'Exempt Write', 'Exempt Listen']
    
    unmelted = temp_df.columns.difference(cols_to_melt)
    # print(unmelted)
    temp_df = temp_df.melt(id_vars=unmelted, value_vars=cols_to_melt, var_name='Exemption', value_name='Value')

    # Fill NaNs with placeholders to prevent issues with groupby
    temp_df = temp_df.fillna("__NA__")
    # Filter the melted DataFrame to only keep rows with "Yes" or keep blank rows
    # Group by the unmelted columns and concatenate 'Exemption' where 'Value' == 'Yes'
    temp_df = temp_df.groupby(list(unmelted)).apply(
        lambda group: pd.Series({
            **group.iloc[0][list(unmelted)],  # Keep the unmelted columns
            'Exemptions': ', '.join(group.loc[group['Value'] == 'Yes', 'Exemption'])
        })
    ).reset_index(drop=True)

    # Replace placeholders back with NaNs
    temp_df.replace("__NA__", np.nan, inplace=True)

    # Replace empty strings with NaN to indicate no exemptions
    temp_df['Exemptions'] = temp_df['Exemptions'].replace('', np.nan)
    print("EXEMPTION COLUMNS MELTED")
    temp_path = elpa_cards_folder / "EL Testing Cards.xlsx"
    
    assets_folder = folders["assets_folder"]
    word_template_path = assets_folder / "ELPA Testing Card.docx"
    print("STARTING THE WRITE PROCESS FOR SHEETS")
    try:
        with pd.ExcelWriter(temp_path) as writer:
            elementary_df = temp_df[temp_df["School"].str.contains("Elementary|Academy", case=False, na=False)]
            if not elementary_df.empty:
                elementary_df = elementary_df.rename(columns={"Home_room":"Teacher"})
                print(f"Columns after renaming: {elementary_df.columns}")
                elementary_df.sort_values(by=["School","Teacher"], inplace=True)
                existing_columns = [col for col in elpa_testing_card_col_elementary if col in elementary_df.columns]
                elementary_df = elementary_df[existing_columns]
                # print(elementary_df.columns)
                # create_doc(word_template_path,elpa_cards_folder,elementary_df,"Elementary School")
                elementary_df.to_excel(writer, sheet_name="Elementary Schools", index=False)
                print("ELEMENTARY SHEET CREATED")
                
                

            middle_df = temp_df[temp_df["School"].str.contains("Middle", na=False)]
            if not middle_df.empty:
                middle_df = middle_df.drop(columns=[col for col in ['Home_room'] if col in middle_df.columns])
                middle_df = pd.merge(middle_df, ela_class_df, on="student_number", how="left")
                middle_df.sort_values(by=["Teacher","Expression"], inplace=True)
                middle_df = middle_df[elpa_testing_card_col_middle]
                # create_doc(word_template_path,elpa_cards_folder,middle_df,"High School")
                # print(middle_df.columns)
                middle_df.to_excel(writer, sheet_name="Middle Schools", index=False)
                print("MIDDLE SHEET CREATED")
                
                

            high_df = temp_df[~temp_df["School"].str.contains("Elementary|Middle|Academy", na=False)]
            if not high_df.empty:
                high_df = high_df.drop(columns=[col for col in ['Home_room'] if col in high_df.columns])
                high_df.sort_values(by=["School","LastFirst"], inplace=True)
                high_df = high_df[elpa_testing_card_col_high]
                # print(high_df.columns)
                #  create_doc(word_template_path,elpa_cards_folder,high_df,"Middle School")
                high_df.to_excel(writer, sheet_name="High Schools", index=False)
                print("HIGH SHEET CREATED")
                
                
            
            # temp_df.to_excel(writer,sheet_name="All ELs",index=False)
            print("ALL ELS SHEET CREATED")
    except Exception as e:
        print(f"Error executing script: {e}")
        return (False, "")

    adjust_column_widths(temp_path)
 
    subprocess.run(["start", "excel", str(temp_path)], shell=True)

    # # Process each record and generate Word documents

def create_doc(template_path, output_folder, df, name):
    print(f"Creating ELPA testing cards for {name}")  # Ensure this line is correctly formatted

    # Create the final document that will contain all the records
    final_doc = Document()

    # Loop through each record in the DataFrame
    for i, record in enumerate(df.to_dict(orient="records")):
        print(f"Processing record {i + 1}")
        record = {key: (value if pd.notna(value) else 'Not Available') for key, value in record.items()}

        # Initialize the template each time and render it with the record
        doc = DocxTemplate(template_path)
        doc.render(record)

        # Save the rendered content temporarily
        temp_path = output_folder / f"temp_{i}.docx"
        doc.save(temp_path)

        # Load the rendered document and append to final document
        rendered_doc = Document(temp_path)
        
        # Append the content of the rendered document into the final document
        for element in rendered_doc.element.body:
            final_doc.element.body.append(element)

    # Apply final formatting to the entire combined document (e.g., left-align paragraphs)
    for paragraph in final_doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the final combined document
    final_output_path = output_folder / f"{name}_ELPA_Testing_Cards_Combined.docx"
    final_doc.save(final_output_path)
    print(f"FINAL COMBINED DOCUMENT SAVED at {final_output_path}")